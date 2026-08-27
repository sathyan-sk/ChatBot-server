# docx_parser_provider.py
"""Requires python-docx: add "python-docx>=1.1.0" to pyproject.toml dependencies."""

import io

from docx import Document

from domain.provider_interfaces.parser_provider import (
    ParsedDocument,
    ParsedSection,
    ParserProviderInterface,
)
from exceptions.domain_exceptions import ProviderError


class DocxParserProvider(ParserProviderInterface):
    async def parse(self, raw_content: bytes, content_type: str) -> ParsedDocument:
        try:
            document = Document(io.BytesIO(raw_content))
            content = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        except Exception as exc:
            raise ProviderError(f"Failed to parse DOCX: {exc}") from exc
        return ParsedDocument(sections=[ParsedSection(title=None, content=content, order=0)])

    def supports(self, content_type: str) -> bool:
        return (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
